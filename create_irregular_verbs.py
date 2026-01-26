"""
Script to create a Word document with irregular verbs in a table format.
Columns: Base (Present), Past, Past Participle
"""

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys

    print("python-docx library not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# Common irregular verbs with their forms
# Format: (base, past, past_participle)
irregular_verbs = [
    ("arise", "arose", "arisen"),
    ("awake", "awoke", "awoken"),
    ("be", "was/were", "been"),
    ("bear", "bore", "borne"),
    ("beat", "beat", "beaten"),
    ("become", "became", "become"),
    ("begin", "began", "begun"),
    ("bend", "bent", "bent"),
    ("bet", "bet", "bet"),
    ("bind", "bound", "bound"),
    ("bite", "bit", "bitten"),
    ("bleed", "bled", "bled"),
    ("blow", "blew", "blown"),
    ("break", "broke", "broken"),
    ("breed", "bred", "bred"),
    ("bring", "brought", "brought"),
    ("broadcast", "broadcast", "broadcast"),
    ("build", "built", "built"),
    ("burn", "burnt/burned", "burnt/burned"),
    ("burst", "burst", "burst"),
    ("buy", "bought", "bought"),
    ("cast", "cast", "cast"),
    ("catch", "caught", "caught"),
    ("choose", "chose", "chosen"),
    ("cling", "clung", "clung"),
    ("come", "came", "come"),
    ("cost", "cost", "cost"),
    ("creep", "crept", "crept"),
    ("cut", "cut", "cut"),
    ("deal", "dealt", "dealt"),
    ("dig", "dug", "dug"),
    ("do", "did", "done"),
    ("draw", "drew", "drawn"),
    ("dream", "dreamt/dreamed", "dreamt/dreamed"),
    ("drink", "drank", "drunk"),
    ("drive", "drove", "driven"),
    ("eat", "ate", "eaten"),
    ("fall", "fell", "fallen"),
    ("feed", "fed", "fed"),
    ("feel", "felt", "felt"),
    ("fight", "fought", "fought"),
    ("find", "found", "found"),
    ("fit", "fit/fitted", "fit/fitted"),
    ("fly", "flew", "flown"),
    ("forbid", "forbade", "forbidden"),
    ("forget", "forgot", "forgotten"),
    ("forgive", "forgave", "forgiven"),
    ("freeze", "froze", "frozen"),
    ("get", "got", "got/gotten"),
    ("give", "gave", "given"),
    ("go", "went", "gone"),
    ("grow", "grew", "grown"),
    ("hang", "hung", "hung"),
    ("have", "had", "had"),
    ("hear", "heard", "heard"),
    ("hide", "hid", "hidden"),
    ("hit", "hit", "hit"),
    ("hold", "held", "held"),
    ("hurt", "hurt", "hurt"),
    ("keep", "kept", "kept"),
    ("kneel", "knelt/kneeled", "knelt/kneeled"),
    ("know", "knew", "known"),
    ("lay", "laid", "laid"),
    ("lead", "led", "led"),
    ("lean", "leant/leaned", "leant/leaned"),
    ("leap", "leapt/leaped", "leapt/leaped"),
    ("learn", "learnt/learned", "learnt/learned"),
    ("leave", "left", "left"),
    ("lend", "lent", "lent"),
    ("let", "let", "let"),
    ("lie", "lay", "lain"),
    ("light", "lit/lighted", "lit/lighted"),
    ("lose", "lost", "lost"),
    ("make", "made", "made"),
    ("mean", "meant", "meant"),
    ("meet", "met", "met"),
    ("mistake", "mistook", "mistaken"),
    ("pay", "paid", "paid"),
    ("prove", "proved", "proven/proved"),
    ("put", "put", "put"),
    ("read", "read", "read"),
    ("ride", "rode", "ridden"),
    ("ring", "rang", "rung"),
    ("rise", "rose", "risen"),
    ("run", "ran", "run"),
    ("say", "said", "said"),
    ("see", "saw", "seen"),
    ("seek", "sought", "sought"),
    ("sell", "sold", "sold"),
    ("send", "sent", "sent"),
    ("set", "set", "set"),
    ("shake", "shook", "shaken"),
    ("shine", "shone", "shone"),
    ("shoot", "shot", "shot"),
    ("show", "showed", "shown/showed"),
    ("shrink", "shrank", "shrunk"),
    ("shut", "shut", "shut"),
    ("sing", "sang", "sung"),
    ("sink", "sank", "sunk"),
    ("sit", "sat", "sat"),
    ("sleep", "slept", "slept"),
    ("slide", "slid", "slid"),
    ("speak", "spoke", "spoken"),
    ("spend", "spent", "spent"),
    ("spin", "spun", "spun"),
    ("split", "split", "split"),
    ("spread", "spread", "spread"),
    ("stand", "stood", "stood"),
    ("steal", "stole", "stolen"),
    ("stick", "stuck", "stuck"),
    ("sting", "stung", "stung"),
    ("strike", "struck", "struck"),
    ("swear", "swore", "sworn"),
    ("sweep", "swept", "swept"),
    ("swim", "swam", "swum"),
    ("swing", "swung", "swung"),
    ("take", "took", "taken"),
    ("teach", "taught", "taught"),
    ("tear", "tore", "torn"),
    ("tell", "told", "told"),
    ("think", "thought", "thought"),
    ("throw", "threw", "thrown"),
    ("understand", "understood", "understood"),
    ("wake", "woke", "woken"),
    ("wear", "wore", "worn"),
    ("weave", "wove", "woven"),
    ("weep", "wept", "wept"),
    ("win", "won", "won"),
    ("wind", "wound", "wound"),
    ("withdraw", "withdrew", "withdrawn"),
    ("withstand", "withstood", "withstood"),
    ("wring", "wrung", "wrung"),
    ("write", "wrote", "written"),
]

# Remove any accidental duplicates and sort alphabetically by base form
seen = set()
unique_irregulars = []
for base, past, pp in irregular_verbs:
    if base not in seen:
        seen.add(base)
        unique_irregulars.append((base, past, pp))

unique_irregulars.sort(key=lambda x: x[0])

# Create Word document
doc = Document()

# Add title
title = doc.add_heading("Irregular Verbs", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph("Base (Present) | Past | Past Participle")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(12)
subtitle_run.font.bold = True

# Blank line
doc.add_paragraph()

# Create table (header + rows)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"

header_cells = table.rows[0].cells
header_cells[0].text = "Base (Present)"
header_cells[1].text = "Past"
header_cells[2].text = "Past Participle"

for cell in header_cells:
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add each irregular verb
for base, past, pp in unique_irregulars:
    row_cells = table.add_row().cells
    row_cells[0].text = base
    row_cells[1].text = past
    row_cells[2].text = pp

    for cell in row_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save directly to the existing irregular_verbs.docx in this folder
output_file = "irregular_verbs.docx"
doc.save(output_file)

print(f"Document created successfully: {output_file}")
print(f"Total irregular verbs: {len(unique_irregulars)}")

